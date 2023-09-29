import datetime
from django.contrib import messages
from django.http import HttpResponse
from django.shortcuts import render, redirect
from shop.models import Item
from docx import Document


def Main(request):
    data = Item.objects.all()
    print(data)
    context = {"data": data}
    return render(request, 'index.html', context)

def ExportWord(request):
    if request.method == 'POST':
       selected_item_ids = request.POST.getlist('selected_items')
       if selected_item_ids:
         selected_items = Item.objects.filter(id__in=selected_item_ids)

         document = Document()
         document.add_heading('Товары', 0)

         table = document.add_table(rows=1, cols=3)
         table.autofit = True
         table.style = "Table Grid"

         table.rows[0].cells[0].text = "ID товара"
         table.rows[0].cells[1].text = "Имя товара"
         table.rows[0].cells[2].text = "Описание"

         for item in selected_items:
             row = table.add_row().cells
             row[0].text = str(item.itemid)
             row[1].text = str(item.itemname)
             row[2].text = str(item.description)

         content = document.add_paragraph("Имя покупателя __________. Вы покупаете:")
         for item in selected_items:
             itemname = content.add_run(" " + item.itemname + ",")
             itemname.bold = True

         response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
         response['Content-Disposition'] = 'attachment; filename=Items_' + \
                 str(datetime.datetime.now()) + '.docx'
         document.save(response)
         return response
       else:
         messages.error(request, "Выберите хотя бы один товар для экспорта")
         return redirect('index-page')